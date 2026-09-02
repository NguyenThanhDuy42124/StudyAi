"""
Parse DOCX → plain text dùng python-docx.
Xử lý: paragraphs, tables, headings.
"""
from docx import Document
from loguru import logger


def parse_docx(file_path: str) -> str:
    """
    Extract toàn bộ text từ file DOCX.
    Bao gồm: paragraphs, headings, table cells.
    
    Args:
        file_path: Đường dẫn đến file DOCX
    
    Returns:
        Text đã extract
    """
    try:
        doc = Document(file_path)
        parts: list[str] = []
        
        # Extract paragraphs (bao gồm headings)
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Thêm prefix cho headings để giữ cấu trúc
                if para.style.name.startswith("Heading"):
                    parts.append(f"\n## {text}\n")
                else:
                    parts.append(text)
        
        # Extract text từ tables
        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    parts.append(" | ".join(row_texts))
        
        result = "\n".join(parts)
        logger.info(f"DOCX parse OK: {len(doc.paragraphs)} paragraphs, {len(result)} ký tự")
        return result
        
    except Exception as e:
        logger.error(f"Lỗi parse DOCX {file_path}: {e}")
        raise
